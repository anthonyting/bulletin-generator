import docx as d
from .. import inputs

def _setDate(paragraphs: list, year: int, month: int, day: int):
    for paragraph in paragraphs:
        paragraphText = paragraph.text
        if ("*YEAR*" in paragraph.text and "*MONTH*" in paragraphText and "*DAY*" in paragraphText):
            for run in paragraph.runs:
                if (run.text == "*YEAR*"):
                    run.text = str(year)
                    continue
                if (run.text == "*MONTH*"):
                    run.text = str(month)
                    continue
                if (run.text == "*DAY*"):
                    run.text = str(day)
                    break

def _setTitle(table, text: str):
    title = table.rows[0].cells[0].paragraphs[0]
    if (title.text == "*FR_TITLE*"):
        title.runs[0].text = text

def makePage(document: d.Document, I: inputs.Inputs):
    _setDate(document.paragraphs, I.year, I.month, I.day)
    _setTitle(document.tables[9], I.front_page_title)
