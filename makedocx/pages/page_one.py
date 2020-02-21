import docx as d
from .. import inputs

def _insertOrderOfWorshipTitle(title: str, header: d.section.Section.first_page_header):
    headerRuns = header.paragraphs[0].runs
    for run in headerRuns:
        if (run.text == "*****"): # identifier for header
            run.text = ""
            run.add_text(title)
            return True

def _insertOrderOfWorship(runs: list, order_of_worship: dict):

    foundKey = False
    for run in runs:
        text = run.text
        text = text.replace('’',"'")

        if (not foundKey and text.isascii()):
        # get text to insert based on the key found at each row
            foundKey = True
            text = text.replace(" ","_")

            # if there is a keyerror here, that probably means the word 
            # document xml is incorrect and a run needs to be fixed
            chineseToInsert = order_of_worship[text.lower()][0].strip() # index 0 is chinese
            englishToInsert = order_of_worship[text.lower()][1].strip() # index 1 is english

            if (englishToInsert):
                englishToInsert = " " + englishToInsert

            if (text.lower() == "sermon"):
                chineseToInsert = "【" + chineseToInsert + "】"

        if ("**" == text): # identifier for chinese text
            run.text = ""
            run.add_tab()
            run.add_text(chineseToInsert)
            continue
        if ("***" == text): # identifier for english text
            run.text = ""
            run.add_text(englishToInsert)
            run.font.size = d.shared.Pt(9)
            break

def _formatParagraph(paragraph: d.text.paragraph.Paragraph, maxTabStop: float) -> bool:
    if ("*****" not in paragraph.text):
        return False

    firstStop = 0.6 * maxTabStop + 1/40 * (len(paragraph.text.replace(" ","")) - 6)  # a nice position around 3/4 of the document

    tabStops = paragraph.paragraph_format.tab_stops
    tabStops.add_tab_stop(d.shared.Inches(firstStop), d.enum.text.WD_TAB_ALIGNMENT.CENTER, d.enum.text.WD_TAB_LEADER.DOTS) # first tab stop

    return True

def makePage(document: d.Document, I: inputs.Inputs):
    _insertOrderOfWorshipTitle(I.order_of_worship_title, document.sections[0].first_page_header)

    for table in document.tables[0:4]:
        for row in table.rows: # format paragraph, then insert text
            paragraph = row.cells[0].paragraphs[0] # one cell per row with one line of text in one paragraph

            if (_formatParagraph(paragraph, I._advanced["maxTabStop"])):# indicates this row does not have a section to replace 
                _insertOrderOfWorship(paragraph.runs, I.order_of_worship)

    return True