import docx as d

import requests
from bs4 import BeautifulSoup
import time

from datetime import datetime, timedelta
import unicodedata

from .. import inputs
from makedocx.pages.docx_exceptions import *
from makedocx.pages.page_two import _deleteParagraph

def _replaceDate(paragraph: d.text.paragraph, monthIdentifier: str, dayIdentifer: str, month: int, day: int):
    if (monthIdentifier in paragraph.text and dayIdentifer in paragraph.text):
        for run in paragraph.runs:
            if (run.text == monthIdentifier):
                run.text = str(month)
                continue
            if (run.text == dayIdentifer):
                run.text = str(day)
                return True
    return False

def _fillSchedules(table, I: inputs.Inputs):
    weekLater = I.date + timedelta(days=7)
    weekLaterMonth = datetime.strftime(weekLater, "%m").lstrip("0")
    weekLaterDay = datetime.strftime(weekLater, "%d").lstrip("0")

    columnIndex = 0
    rowIndex = 0
    columnIter = iter(table.columns)
    next(columnIter)
    for column in columnIter:
        cellIter = iter(column.cells)
        next(cellIter)
        for cell in cellIter:
            paragraph = cell.paragraphs[0]
            if (_replaceDate(paragraph, "*MONTH*", "*DAY*", I.month, I.day)):
                continue
            elif (_replaceDate(paragraph, "*NMONTH*", "*NDAY*", weekLaterMonth, weekLaterDay)):
                continue
            else:
                paragraph.runs[0].text = I.schedules[rowIndex][columnIndex]
            rowIndex += 1
        rowIndex = 0
        columnIndex += 1

def _fillScriptures(document: d.Document, I: inputs.Inputs, english: str, chinese:str):

    for paragraph in document.paragraphs:
        if ("*SC_CH*" in paragraph.text):
            for run in paragraph.runs:
                if (run.text == "*SC_CH*"):
                    run.text = chinese
                    break
        elif ("*SC_EN*" in paragraph.text):
            for run in paragraph.runs:
                if (run.text == "*SC_EN*"):
                    run.text = english
                    break
    
    table = document.tables[7]
    runs = table.rows[0].cells[0].paragraphs[0].runs
    for run in runs:
        if (run.text == "**"):
            run.text = f"{I.scriptures['book_ch']}"
        elif (run.text == "***"):
            run.text = f"{I.scriptures['book']} {I.scriptures['chapter']}:{I.scriptures['verses']}"

def _getBibleVerses(book: str, chapter: int, verses: str, version: str):
    url = f"https://www.biblegateway.com/passage/?search={book}+{chapter}%3A{verses}&version={version}"
    r = _scrape(url)
    if (type(r) != requests.models.Response):
        raise r # error in request
    soup = BeautifulSoup(r.text, "html.parser")

    if (version == "CUV"):
        poetry = soup.find("div", {"class": "version-CUV result-text-style-normal text-html"})
    elif (version == "NIV"):
        poetry = soup.find("div", {"class": "version-NIV result-text-style-normal text-html"})
    else:
        raise VerseNotFound("Invalid bible version: use NIV or CUV")

    textSearch = "p"
    if (not poetry):
        raise VerseNotFound("Check bible verse/chapter")

    text = poetry(textSearch)

    # removes leading spaces, leading numbers, and extra whitespace in between text
    return [unicodedata.normalize("NFKD", " ".join(verse.text.strip().lstrip("0123456789 ").split())) for verse in text]

def _scrape(url: str):
    attempts = 0
    while (attempts < 10):
        try:
            r = requests.get(url)
            r.raise_for_status()
            return r
        except requests.exceptions.HTTPError as e:
            return e
        except requests.exceptions.RequestException as e:
            error = e
            attempts += 1
            time.sleep(1)
    return error

def _fillPrayers(paragraphs: list, I: inputs.Inputs):
    foundFirst = False
    index = 0
    for paragraph in paragraphs:
        if (not foundFirst and paragraph.text == "*PR_START*"):
            foundFirst = True

        if (paragraph.text == "*PR_START*"):
            _deleteParagraph(paragraph)
        elif (paragraph.text == "*PR_END*"):
            _deleteParagraph(paragraph)
            return  
        elif (foundFirst):
            if (index < I.prayer_count):
                # each paragraph is set to have two runs
                paragraph.runs[0].text = f"{I.prayers[index]}"
            else:
                _deleteParagraph(paragraph)
            index += 1

def _fillLastBox(paragraphs: list, text: str):
    for paragraph in paragraphs:
        if ("**LAST_BOX**" in paragraph.text):
            paragraph.runs[0].text = text
            return

def makePage(document: d.Document, I: inputs.Inputs):
    _fillPrayers(document.paragraphs, I)
    _fillSchedules(document.tables[6], I)
    _fillLastBox(document.tables[8].rows[0].cells[0].paragraphs, I.last_box) # one box table

    englishVerses = _getBibleVerses(I.scriptures["book"], I.scriptures["chapter"], I.scriptures["verses"], "NIV")
    chineseVerses = _getBibleVerses(I.scriptures["book"], I.scriptures["chapter"], I.scriptures["verses"], "CUV")

    _fillScriptures(document, I, " ".join(englishVerses), " ".join(chineseVerses).replace(" ",""))